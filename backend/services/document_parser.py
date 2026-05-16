"""
Document Parser — Extracts text from PDF (text + scanned), DOCX files.
Uses PyMuPDF for text PDFs, pytesseract for scanned PDFs, python-docx for DOCX.
pdfplumber is used specifically for table extraction.
"""

import io
import os
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from PIL import Image

# Attempt pytesseract import — OCR is optional
try:
    import pytesseract
    # Auto-detect Tesseract on Windows
    _win_tesseract = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.name == "nt" and os.path.exists(_win_tesseract):
        pytesseract.pytesseract.tesseract_cmd = _win_tesseract
    # Verify the binary actually works
    pytesseract.get_tesseract_version()
    HAS_TESSERACT = True
except (ImportError, Exception):
    HAS_TESSERACT = False


def parse_uploaded_file(uploaded_file) -> dict:
    """
    Main entry point. Accepts a Streamlit UploadedFile (or any file-like
    object with .name attribute) and returns extracted content.

    Returns:
        {
            "text": str,           # Full extracted text
            "tables": list[list],  # Extracted tables (list of rows)
            "pages": int,          # Number of pages (PDFs) or paragraphs (DOCX)
            "method": str,         # "pymupdf" | "ocr" | "docx"
            "filename": str
        }
    """
    filename = getattr(uploaded_file, "name", "unknown")
    ext = Path(filename).suffix.lower()
    raw_bytes = uploaded_file.read()
    # Reset stream position so it can be re-read if needed
    if hasattr(uploaded_file, "seek"):
        uploaded_file.seek(0)

    if ext == ".docx":
        return _parse_docx(raw_bytes, filename)
    elif ext == ".pdf":
        return _parse_pdf(raw_bytes, filename)
    elif ext == ".txt":
        text = raw_bytes.decode("utf-8", errors="replace")
        lines = text.strip().splitlines()
        return {
            "text": text,
            "tables": [],
            "pages": max(1, len(lines) // 50),
            "method": "plaintext",
            "filename": filename,
        }
    else:
        return {
            "text": "",
            "tables": [],
            "pages": 0,
            "method": "unsupported",
            "filename": filename,
            "error": f"Unsupported file type: {ext}",
        }


# ── PDF Parsing ────────────────────────────────────────────────────────────


def _parse_pdf(raw_bytes: bytes, filename: str) -> dict:
    """Parse a PDF: try text extraction first, fall back to OCR if mostly empty."""
    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    pages = len(doc)

    # 1. Try text-based extraction with PyMuPDF
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text("text"))
    full_text = "\n".join(text_parts).strip()
    doc.close()

    method = "pymupdf"

    # 2. If text is mostly empty (< 100 chars per page on avg), try OCR
    avg_chars = len(full_text) / max(pages, 1)
    if avg_chars < 50 and HAS_TESSERACT:
        full_text = _ocr_pdf(raw_bytes)
        method = "ocr"
    elif avg_chars < 50:
        full_text = (
            "[Scanned PDF detected but Tesseract OCR is not installed. "
            "Install Tesseract to enable OCR support.]\n\n" + full_text
        )
        method = "pymupdf (limited — OCR unavailable)"

    # 3. Extract tables with pdfplumber
    tables = _extract_tables(raw_bytes)

    return {
        "text": full_text,
        "tables": tables,
        "pages": pages,
        "method": method,
        "filename": filename,
    }


def _ocr_pdf(raw_bytes: bytes) -> str:
    """OCR a scanned PDF page-by-page using pytesseract."""
    doc = fitz.open(stream=raw_bytes, filetype="pdf")
    ocr_texts = []
    for page_num, page in enumerate(doc):
        # Render page to image at 300 DPI for good OCR quality
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        text = pytesseract.image_to_string(img, lang="eng")
        ocr_texts.append(f"--- Page {page_num + 1} ---\n{text}")
    doc.close()
    return "\n\n".join(ocr_texts)


def _extract_tables(raw_bytes: bytes) -> list:
    """Extract tables from PDF using pdfplumber."""
    tables = []
    try:
        with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                if page_tables:
                    for tbl in page_tables:
                        # Filter out empty rows
                        cleaned = [row for row in tbl if any(cell for cell in row)]
                        if cleaned:
                            tables.append(cleaned)
    except Exception:
        pass  # Table extraction is best-effort
    return tables


# ── DOCX Parsing ───────────────────────────────────────────────────────────


def _parse_docx(raw_bytes: bytes, filename: str) -> dict:
    """Parse a DOCX file using python-docx."""
    doc = Document(io.BytesIO(raw_bytes))

    # Extract paragraphs with heading detection
    text_parts = []
    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "Heading" in style:
            text_parts.append(f"\n{'#' * _heading_level(style)} {text}\n")
        else:
            text_parts.append(text)

    # Extract tables
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)

    full_text = "\n".join(text_parts)
    return {
        "text": full_text,
        "tables": tables,
        "pages": len(doc.paragraphs),
        "method": "docx",
        "filename": filename,
    }


def _heading_level(style_name: str) -> int:
    """Extract heading level from Word style name like 'Heading 2'."""
    for ch in style_name:
        if ch.isdigit():
            return min(int(ch), 4)
    return 2
